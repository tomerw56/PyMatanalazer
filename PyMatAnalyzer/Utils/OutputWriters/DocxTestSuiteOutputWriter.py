
from PyMatAnalyzer.Tests.BaseClasses.IOutputWriter import IOutputWriter
from typing import Dict
import os
from docx import Document
from docx.shared import Inches
import pandas as pd

class DocxTestSuiteOutputWriter(IOutputWriter):
    def __init__(self,logger,file_full_name:str,header_str:str):
        self.__logger=logger
        self.__file_full_name=file_full_name
        self.__document = Document()
        self.__document.add_heading(header_str, 0)
        self.__supported_keys={'paragraph_str':'paragraph',
                               'paragraph_header':'paragrpah header',
                               'image_file':'file name for the image',
                               'table_data_frame':'data frame with data to write'}
        self.__unnamed_headers_counter=0
    def WriteTestOutput(self, **kwargs)->bool:
        try:
            # write Header
            if('paragraph_header' not in kwargs.keys()):
                self.__unnamed_headers_counter+=1
                header_str=format(f"Header_{self.__unnamed_headers_counter}")
                self.__document.add_heading(header_str, level = 1)
                self.__unnamed_headers_counter+=1
            else:
                self.__document.add_heading(kwargs['paragraph_header'], level=1)
            #add  str
            if ('paragraph_str' in kwargs.keys()):
                self.__document.add_paragraph(kwargs['paragraph_str'])
            #write image
            if ('image_file' in kwargs.keys()):
                self.__document.add_picture(kwargs['image_file'])
            # write table
            if ('table_data_frame' in kwargs.keys()):
                if(type(kwargs['table_data_frame']) is not pd.DataFrame):
                    self.__logger.warning('recived data not suitible for table -expecting data frame')
                else:
                    data=(pd.DataFrame)(kwargs['table_data_frame'])
                    number_of_cols=len(data.keys())
                    number_of_rows=data.shape[0]
                    table = self.__document.add_table(rows=1, cols=number_of_cols,style='Light List')
                    hdr_cells = table.rows[0].cells
                    column_counter=0
                    for key in data.keys():
                        hdr_cells[column_counter].text = key
                        column_counter+=1
                    items_list=data.to_dict('records')
                    for item in items_list:
                        row_cells = table.add_row().cells
                        row_counter=0
                        for key in data.keys():
                            row_cells[row_counter].text = str(item[key])
                            row_counter += 1

            #break any way
            self.__document.add_page_break()





        except Exception as e:
            self.__logger.error(f'could not write data {e}')
            return False
        return True
    #since we want to document in an entire doc the test suite we allow the saving out side
    def EndTestSuiteWrite(self)->bool:
        try:
            self.__document.save(self.__file_full_name)
            return True
        except Exception as e:
            self.__logger.error(f"save encountered an error {e}")
            return False
    def SupportedKeys(self)->Dict[str,str]:
        return self.__supported_keys

    def EndTestWriting(self):
        pass